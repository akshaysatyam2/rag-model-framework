# # from flask import Flask, request, jsonify, render_template
# # from werkzeug.utils import secure_filename
# # import os
# # import tempfile
# # import docx
# # import PyPDF2
# # from transformers import pipeline
# # from sentence_transformers import SentenceTransformer, CrossEncoder
# # import faiss
# # import numpy as np
# # import torch

# # app = Flask(__name__)

# # UPLOAD_FOLDER = tempfile.gettempdir()
# # ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx'}
# # app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# # device = 0 if torch.cuda.is_available() else 'cpu'

# # # Load models
# # retriever_model = SentenceTransformer('multi-qa-mpnet-base-dot-v1')
# # reranker = CrossEncoder('cross-encoder/ms-marco-MiniLM-L-12-v2', device=device)
# # qa_model = pipeline("question-answering", model="deepset/roberta-base-squad2", device=device)
# # summarizer = pipeline("summarization", model="facebook/bart-large-cnn", device=device)

# # pipeline_components = {}

# # # Helper functions
# # def allowed_file(filename):
# #     return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# # def read_file(filepath):
# #     ext = filepath.rsplit('.', 1)[1].lower()
# #     if ext == 'txt':
# #         with open(filepath, 'r', encoding='utf-8') as f:
# #             return f.read()
# #     elif ext == 'pdf':
# #         with open(filepath, 'rb') as f:
# #             reader = PyPDF2.PdfReader(f)
# #             return ' '.join(page.extract_text() for page in reader.pages)
# #     elif ext == 'docx':
# #         doc = docx.Document(filepath)
# #         return ' '.join(paragraph.text for paragraph in doc.paragraphs)
# #     else:
# #         raise ValueError("Unsupported file format.")

# # def split_text(text, chunk_size=400):
# #     # Split text into meaningful chunks based on sentences
# #     sentences = text.split('. ')
# #     chunks = []
# #     current_chunk = ""

# #     for sentence in sentences:
# #         if len(current_chunk) + len(sentence) < chunk_size:
# #             current_chunk += sentence + ". "
# #         else:
# #             chunks.append(current_chunk.strip())
# #             current_chunk = sentence + ". "

# #     if current_chunk:
# #         chunks.append(current_chunk.strip())

# #     return chunks

# # def build_retrieval_pipeline(text):
# #     chunks = split_text(text)
# #     embeddings = retriever_model.encode(chunks, convert_to_tensor=True, batch_size=8).cpu().numpy().astype('float32')

# #     dimension = embeddings.shape[1]
# #     index = faiss.IndexHNSWFlat(dimension, 32)
# #     index.hnsw.efSearch = 64
# #     index.add(embeddings)

# #     return {
# #         "chunks": chunks,
# #         "index": index
# #     }

# # def improved_rag_pipeline(question, pipeline_data, top_k=15):
# #     # Retrieve relevant chunks
# #     question_embedding = retriever_model.encode([question], convert_to_tensor=True).cpu().numpy().astype('float32')
# #     distances, indices = pipeline_data['index'].search(question_embedding, k=top_k)
# #     retrieved_chunks = [pipeline_data['chunks'][i] for i in indices[0]]

# #     # Rerank chunks based on relevance
# #     rerank_scores = reranker.predict([(question, chunk) for chunk in retrieved_chunks])
# #     ranked_chunks = [chunk for _, chunk in sorted(zip(rerank_scores, retrieved_chunks), reverse=True)]

# #     # Generate answers
# #     context = " ".join(ranked_chunks[:5])  # Use top 5 chunks for context
# #     qa_result = qa_model(question=question, context=context)
# #     direct_answer = qa_result['answer']
# #     enriched_answer = summarizer(context, min_length=25, do_sample=False)[0]['summary_text']

# #     return {
# #         "Direct_Answer": direct_answer,
# #         "Enriched_Contextual_Answer": enriched_answer
# #     }

# # # Routes
# # @app.route('/')
# # def index():
# #     return render_template('index.html')

# # @app.route('/upload', methods=['POST'])
# # def upload_file():
# #     if 'file' not in request.files:
# #         return jsonify(error="No file uploaded."), 400

# #     file = request.files['file']
# #     if file.filename == '':
# #         return jsonify(error="No selected file."), 400

# #     if file and allowed_file(file.filename):
# #         filename = secure_filename(file.filename)
# #         filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
# #         file.save(filepath)

# #         try:
# #             text = read_file(filepath)
# #             global pipeline_components
# #             pipeline_components = build_retrieval_pipeline(text)
# #             return jsonify(message="File uploaded and processed successfully.")
# #         except Exception as e:
# #             return jsonify(error=str(e)), 500
# #     else:
# #         return jsonify(error="Unsupported file type. Only txt, pdf, and docx are allowed."), 400

# # @app.route('/ask', methods=['POST'])
# # def ask_question():
# #     if not pipeline_components:
# #         return jsonify(error="No file uploaded or processed."), 400

# #     data = request.get_json()
# #     question = data.get('query')
# #     if not question:
# #         return jsonify(error="Question cannot be empty."), 400

# #     try:
# #         answer = improved_rag_pipeline(question, pipeline_components)
# #         return jsonify(answer=answer['Direct_Answer'], enriched=answer['Enriched_Contextual_Answer'])
# #     except Exception as e:
# #         return jsonify(error=str(e)), 500

# # if __name__ == '__main__':
# #     app.run(debug=True)


# from flask import Flask, request, jsonify, render_template
# from werkzeug.utils import secure_filename
# import os
# import tempfile
# import docx
# import PyPDF2
# from transformers import pipeline
# from sentence_transformers import SentenceTransformer, CrossEncoder
# import faiss
# import numpy as np
# import torch

# app = Flask(__name__)

# UPLOAD_FOLDER = tempfile.gettempdir()
# ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx'}
# app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# device = 0 if torch.cuda.is_available() else 'cpu'

# # Models
# retriever_model_mpnet = SentenceTransformer('multi-qa-mpnet-base-dot-v1')
# retriever_model_minilm = SentenceTransformer('all-MiniLM-L6-v2')
# reranker = CrossEncoder('cross-encoder/ms-marco-MiniLM-L-12-v2', device=device)
# qa_model = pipeline("question-answering", model="deepset/roberta-base-squad2", device=device)
# summarizer = pipeline("summarization", model="facebook/bart-large-cnn", device=device)

# pipeline_components = {}
# retriever_model = retriever_model_mpnet

# def allowed_file(filename):
#     return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# def read_file(filepath):
#     ext = filepath.rsplit('.', 1)[1].lower()
#     if ext == 'txt':
#         with open(filepath, 'r', encoding='utf-8') as f:
#             return f.read()
#     elif ext == 'pdf':
#         with open(filepath, 'rb') as f:
#             reader = PyPDF2.PdfReader(f)
#             return ' '.join(page.extract_text() for page in reader.pages)
#     elif ext == 'docx':
#         doc = docx.Document(filepath)
#         return ' '.join(paragraph.text for paragraph in doc.paragraphs)
#     else:
#         raise ValueError("Unsupported file format.")

# def split_text(text, chunk_size=150, overlap=50):
#     words = text.split()
#     return [" ".join(words[i:i + chunk_size]) for i in range(0, len(words), chunk_size - overlap)]

# def build_retrieval_pipeline(text, embedder='mpnet'):
#     global retriever_model
#     retriever_model = retriever_model_mpnet if embedder == 'mpnet' else retriever_model_minilm

#     chunks = split_text(text)
#     embeddings = retriever_model.encode(chunks, convert_to_tensor=True, batch_size=8).cpu().numpy().astype('float32')

#     dimension = embeddings.shape[1]
#     index = faiss.IndexHNSWFlat(dimension, 32)
#     index.hnsw.efSearch = 64
#     index.add(embeddings)

#     return {
#         "chunks": chunks,
#         "index": index
#     }

# def improved_rag_pipeline(question, pipeline_data, top_k=15):
#     question_embedding = retriever_model.encode([question], convert_to_tensor=True).cpu().numpy().astype('float32')
#     distances, indices = pipeline_data['index'].search(question_embedding, k=top_k)
#     retrieved_chunks = [pipeline_data['chunks'][i] for i in indices[0]]

#     rerank_scores = reranker.predict([(question, chunk) for chunk in retrieved_chunks])
#     ranked_chunks = [chunk for _, chunk in sorted(zip(rerank_scores, retrieved_chunks), reverse=True)]

#     context = " ".join(ranked_chunks[:5])
#     qa_result = qa_model(question=question, context=context)
#     direct_answer = qa_result['answer']
#     enriched_answer = summarizer(context, min_length=25, do_sample=False)[0]['summary_text']

#     return {
#         "Direct_Answer": direct_answer,
#         "Enriched_Contextual_Answer": enriched_answer
#     }

# @app.route('/')
# def index():
#     return render_template('index.html')

# @app.route('/upload', methods=['POST'])
# def upload_file():
#     if 'file' not in request.files:
#         return jsonify(error="No file uploaded."), 400

#     file = request.files['file']
#     if file.filename == '':
#         return jsonify(error="No selected file."), 400

#     if file and allowed_file(file.filename):
#         filename = secure_filename(file.filename)
#         filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
#         file.save(filepath)

#         try:
#             embedder = request.form.get('embedder', 'mpnet')
#             text = read_file(filepath)
#             global pipeline_components
#             pipeline_components = build_retrieval_pipeline(text, embedder)
#             return jsonify(message="File uploaded and processed successfully.", embedder=embedder)
#         except Exception as e:
#             return jsonify(error=str(e)), 500
#     else:
#         return jsonify(error="Unsupported file type. Only txt, pdf, and docx are allowed."), 400

# @app.route('/ask', methods=['POST'])
# def ask_question():
#     if not pipeline_components:
#         return jsonify(error="No file uploaded or processed."), 400

#     data = request.get_json()
#     question = data.get('query')
#     if not question:
#         return jsonify(error="Question cannot be empty."), 400

#     try:
#         answer = improved_rag_pipeline(question, pipeline_components)
#         return jsonify(answer=answer['Direct_Answer'], enriched=answer['Enriched_Contextual_Answer'])
#     except Exception as e:
#         return jsonify(error=str(e)), 500

# if __name__ == '__main__':
#     app.run(debug=True)


# main.py
from flask import Flask, request, jsonify, render_template
from werkzeug.utils import secure_filename
import os
import tempfile
import docx
import PyPDF2
from transformers import pipeline
from sentence_transformers import SentenceTransformer, CrossEncoder
import faiss
import numpy as np
import torch

# Initialize the Flask application
app = Flask(__name__)

UPLOAD_FOLDER = tempfile.gettempdir()
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

device = 0 if torch.cuda.is_available() else 'cpu'

# Load models
retriever_model_mpnet = SentenceTransformer('multi-qa-mpnet-base-dot-v1')
retriever_model_minilm = SentenceTransformer('all-MiniLM-L6-v2')
reranker = CrossEncoder('cross-encoder/ms-marco-MiniLM-L-12-v2', device=device)
qa_model = pipeline("question-answering", model="deepset/roberta-base-squad2", device=device)
summarizer = pipeline("summarization", model="facebook/bart-large-cnn", device=device)

pipeline_components = {}
retriever_model = retriever_model_mpnet

# Helper functions
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def read_file(filepath):
    ext = filepath.rsplit('.', 1)[1].lower()
    if ext == 'txt':
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    elif ext == 'pdf':
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            return ' '.join(page.extract_text() for page in reader.pages)
    elif ext == 'docx':
        doc = docx.Document(filepath)
        return ' '.join(paragraph.text for paragraph in doc.paragraphs)
    else:
        raise ValueError("Unsupported file format.")

def split_text(text, chunk_size=150, overlap=50):
    words = text.split()
    return [" ".join(words[i:i + chunk_size]) for i in range(0, len(words), chunk_size - overlap)]

def build_retrieval_pipeline(text, embedder='mpnet'):
    global retriever_model
    retriever_model = retriever_model_mpnet if embedder == 'mpnet' else retriever_model_minilm

    chunks = split_text(text)
    embeddings = retriever_model.encode(chunks, convert_to_tensor=True, batch_size=8).cpu().numpy().astype('float32')

    dimension = embeddings.shape[1]
    index = faiss.IndexHNSWFlat(dimension, 32)
    index.hnsw.efSearch = 64
    index.add(embeddings)

    return {
        "chunks": chunks,
        "index": index
    }

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify(error="No file uploaded."), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify(error="No selected file."), 400

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        try:
            embedder = request.form.get('embedder', 'mpnet')
            text = read_file(filepath)
            global pipeline_components
            pipeline_components = build_retrieval_pipeline(text, embedder)
            return jsonify(message="File uploaded and processed successfully.", embedder=embedder)
        except Exception as e:
            return jsonify(error=str(e)), 500
    else:
        return jsonify(error="Unsupported file type. Only txt, pdf, and docx are allowed."), 400

@app.route('/ask', methods=['POST'])
def ask_question():
    if not pipeline_components:
        return jsonify(error="No file uploaded or processed."), 400

    data = request.get_json()
    question = data.get('query')
    if not question:
        return jsonify(error="Question cannot be empty."), 400

    try:
        question_embedding = retriever_model.encode([question], convert_to_tensor=True).cpu().numpy().astype('float32')
        distances, indices = pipeline_components['index'].search(question_embedding, k=15)
        retrieved_chunks = [pipeline_components['chunks'][i] for i in indices[0]]

        rerank_scores = reranker.predict([(question, chunk) for chunk in retrieved_chunks])
        ranked_chunks = [chunk for _, chunk in sorted(zip(rerank_scores, retrieved_chunks), reverse=True)]

        context = " ".join(ranked_chunks[:5])
        qa_result = qa_model(question=question, context=context)
        direct_answer = qa_result['answer']
        enriched_answer = summarizer(context, min_length=25, do_sample=False)[0]['summary_text']

        return jsonify(answer=direct_answer, enriched=enriched_answer)
    except Exception as e:
        return jsonify(error=str(e)), 500

if __name__ == '__main__':
    app.run(debug=True)
